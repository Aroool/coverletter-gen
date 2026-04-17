"""
Builds the .docx cover letter matching the exact format of the user's template:
- Times New Roman throughout
- Name: Bold 28pt centered
- Contact line: Bold centered, separated by " | "
- Page border: single line all 4 sides
- Margins: 0.5" all around
- Body: auto-spaced paragraphs
- Bullets using proper numbering (not unicode characters)
- Closing: Sincerely + bold name
"""

from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


TNR = "Times New Roman"


# ── XML helpers ────────────────────────────────────────────────────────────────

def set_font(run, size_pt: float, bold: bool = False, italic: bool = False, color=None):
    """Apply Times New Roman with given size and style to a run."""
    run.font.name = TNR
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Force east Asia font too
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), TNR)
    rFonts.set(qn("w:eastAsia"), TNR)
    rFonts.set(qn("w:hAnsi"), TNR)
    rFonts.set(qn("w:cs"), TNR)
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def set_paragraph_spacing(para, before_auto=True, after_auto=True,
                          line_rule="auto", line="240", after_pt=6):
    """
    Paragraph spacing.
    line="240" = single (1.0), "276" = 1.15, "288" = 1.2
    after_pt = gap after each paragraph in points (72pt = 1 inch)
    """
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    if before_auto:
        spacing.set(qn("w:beforeAutospacing"), "1")
        spacing.set(qn("w:before"), "100")
    # Explicit after spacing in twips (1pt = 20 twips)
    spacing.set(qn("w:after"), str(after_pt * 20))
    spacing.set(qn("w:line"), line)
    spacing.set(qn("w:lineRule"), line_rule)
    existing = pPr.find(qn("w:spacing"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(spacing)


def add_page_border(doc):
    """Add single-line page border on all 4 sides (matching template)."""
    sectPr = doc.sections[0]._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "24")
        border.set(qn("w:color"), "auto")
        pgBorders.append(border)
    # Remove existing if any
    existing = sectPr.find(qn("w:pgBorders"))
    if existing is not None:
        sectPr.remove(existing)
    sectPr.append(pgBorders)


def add_hyperlink(para, text: str, url: str, bold: bool = False, size_pt: float = 12):
    """Add a hyperlink run to a paragraph."""
    part = para.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    hyperlink.set(qn("w:history"), "1")

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), TNR)
    rFonts.set(qn("w:eastAsia"), TNR)
    rFonts.set(qn("w:hAnsi"), TNR)
    rFonts.set(qn("w:cs"), TNR)
    rPr.append(rFonts)

    if bold:
        b = OxmlElement("w:b")
        bCs = OxmlElement("w:bCs")
        rPr.append(b)
        rPr.append(bCs)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)
    rPr.append(szCs)

    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    para._p.append(hyperlink)


def setup_bullet_numbering(doc):
    """Add bullet list numbering definition to the document."""
    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        # Create numbering part
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        from docx.oxml import parse_xml
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        numbering_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="bullet"/>
      <w:lvlText w:val="&#x2022;"/>
      <w:lvlJc w:val="left"/>
      <w:pPr>
        <w:ind w:left="360" w:hanging="360"/>
      </w:pPr>
      <w:rPr>
        <w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/>
      </w:rPr>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>'''
        numbering_part_obj = Part(
            PackURI("/word/numbering.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
            parse_xml(numbering_xml.encode()),
            doc.part.package
        )
        doc.part.relate_to(numbering_part_obj, RT.NUMBERING)
        return 1

    # Add to existing numbering part
    numbering_elem = numbering_part._element
    # Find max abstractNumId
    abstract_nums = numbering_elem.findall(qn("w:abstractNum"))
    max_id = max((int(a.get(qn("w:abstractNumId"), -1)) for a in abstract_nums), default=-1)
    new_abstract_id = max_id + 1

    nums = numbering_elem.findall(qn("w:num"))
    max_num_id = max((int(n.get(qn("w:numId"), 0)) for n in nums), default=0)
    new_num_id = max_num_id + 1

    abstract_xml = f'''<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:abstractNumId="{new_abstract_id}">
  <w:lvl w:ilvl="0">
    <w:start w:val="1"/>
    <w:numFmt w:val="bullet"/>
    <w:lvlText w:val="&#x2022;"/>
    <w:lvlJc w:val="left"/>
    <w:pPr>
      <w:ind w:left="360" w:hanging="360"/>
    </w:pPr>
    <w:rPr>
      <w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/>
    </w:rPr>
  </w:lvl>
</w:abstractNum>'''

    num_xml = f'''<w:num xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:numId="{new_num_id}">
  <w:abstractNumId w:val="{new_abstract_id}"/>
</w:num>'''

    from docx.oxml import parse_xml
    abstract_elem = parse_xml(abstract_xml.encode())
    num_elem = parse_xml(num_xml.encode())

    # Insert before existing nums, append num at end
    first_num = numbering_elem.find(qn("w:num"))
    if first_num is not None:
        numbering_elem.insert(list(numbering_elem).index(first_num), abstract_elem)
    else:
        numbering_elem.append(abstract_elem)
    numbering_elem.append(num_elem)

    return new_num_id


def add_bullet_paragraph(doc, text: str, num_id: int, size_pt: float = 11):
    """Add a bullet list paragraph."""
    para = doc.add_paragraph()
    set_paragraph_spacing(para)

    pPr = para._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId_elem = OxmlElement("w:numId")
    numId_elem.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId_elem)
    pPr.insert(0, numPr)

    run = para.add_run(text)
    set_font(run, size_pt)
    return para


# ── Main builder ───────────────────────────────────────────────────────────────

def build_docx(
    profile: dict,
    content: dict,
    company_name: str,
    role_title: str,
    today: date,
    output_path: Path,
):
    doc = Document()

    # ── Page setup: US Letter (8.5 x 11in), 0.5" margins ────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.5)   # 7,772,400 EMU
    section.page_height = Inches(11)    # 10,058,400 EMU
    half_inch = Inches(0.5)
    section.top_margin    = half_inch
    section.bottom_margin = half_inch
    section.left_margin   = half_inch
    section.right_margin  = half_inch

    # Remove default paragraph spacing from Normal style
    normal_style = doc.styles["Normal"]
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(0)

    # ── Page border ───────────────────────────────────────────────────────────
    add_page_border(doc)

    # ── Setup bullet numbering ────────────────────────────────────────────────
    num_id = setup_bullet_numbering(doc)

    # ── Header: Name ─────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(name_para)
    name_run = name_para.add_run(profile.get("name", ""))
    set_font(name_run, size_pt=28, bold=True)

    # ── Header: Contact line ─────────────────────────────────────────────────
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(contact_para)

    email = profile.get("email", "")
    phone = profile.get("phone", "")
    linkedin_url = profile.get("linkedin_url", "")
    github_url = profile.get("github_url", "")
    portfolio_url = profile.get("portfolio_url", "")
    location = profile.get("location", "")

    def sep(para):
        r = para.add_run(" | ")
        set_font(r, size_pt=11, bold=True)

    # Email (hyperlink)
    if email:
        add_hyperlink(contact_para, email, f"mailto:{email}", bold=True, size_pt=11)

    # Phone
    if phone:
        sep(contact_para)
        ph_run = contact_para.add_run(phone)
        set_font(ph_run, size_pt=11, bold=True)

    # LinkedIn (hyperlink)
    if linkedin_url:
        sep(contact_para)
        add_hyperlink(contact_para, "LinkedIn", linkedin_url, bold=True, size_pt=11)

    # GitHub (hyperlink)
    if github_url:
        sep(contact_para)
        add_hyperlink(contact_para, "GitHub", github_url, bold=True, size_pt=11)

    # Portfolio (hyperlink)
    if portfolio_url:
        sep(contact_para)
        add_hyperlink(contact_para, "Portfolio", portfolio_url, bold=True, size_pt=11)

    # Location
    if location:
        sep(contact_para)
        loc_run = contact_para.add_run(location)
        set_font(loc_run, size_pt=11, bold=True)

    # ── Date ─────────────────────────────────────────────────────────────────
    date_para = doc.add_paragraph()
    set_paragraph_spacing(date_para)
    date_run = date_para.add_run(today.strftime("%B %d, %Y"))
    set_font(date_run, size_pt=11)

    # ── Recipient block ───────────────────────────────────────────────────────
    hm_name = content.get("hiring_manager_name", "Hiring Manager")
    hm_title = content.get("hiring_manager_title", "")

    hm_para = doc.add_paragraph()
    set_paragraph_spacing(hm_para)
    hm_run = hm_para.add_run(hm_name)
    set_font(hm_run, size_pt=11, bold=True)

    if hm_title:
        hm_title_para = doc.add_paragraph()
        set_paragraph_spacing(hm_title_para)
        hm_title_run = hm_title_para.add_run(hm_title)
        set_font(hm_title_run, size_pt=11)

    company_para = doc.add_paragraph()
    set_paragraph_spacing(company_para)
    company_run = company_para.add_run(company_name)
    set_font(company_run, size_pt=11)

    # ── Opening hook ─────────────────────────────────────────────────────────
    opening_para = doc.add_paragraph()
    opening_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(opening_para, line="276", after_pt=6)
    opening_run = opening_para.add_run(content.get("opening_hook", ""))
    set_font(opening_run, size_pt=11)

    # ── Body paragraphs ───────────────────────────────────────────────────────
    for body_text in content.get("body_paragraphs", []):
        if not body_text or not body_text.strip():
            continue
        body_para = doc.add_paragraph()
        body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_spacing(body_para, line="276", after_pt=6)

        # Handle inline project references — find bold segments (project names)
        # Convention: if content includes **ProjectName**, make it bold
        import re
        parts = re.split(r'\*\*(.+?)\*\*', body_text)
        for i, part in enumerate(parts):
            if not part:
                continue
            run = body_para.add_run(part)
            set_font(run, size_pt=11, bold=(i % 2 == 1))

    # ── Value bullets ─────────────────────────────────────────────────────────
    bullets = [b for b in content.get("value_bullets", []) if b and b.strip()]
    if bullets:
        # Intro sentence for bullets
        intro_para = doc.add_paragraph()
        set_paragraph_spacing(intro_para, line="276", after_pt=4)
        intro_run = intro_para.add_run(
            f"What I'd bring to {company_name}:"
        )
        set_font(intro_run, size_pt=11)

        for bullet_text in bullets:
            add_bullet_paragraph(doc, bullet_text, num_id, size_pt=11)

    # ── Closing paragraph ─────────────────────────────────────────────────────
    closing_text = content.get("closing_paragraph", "")
    if closing_text:
        closing_para = doc.add_paragraph()
        closing_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_spacing(closing_para, line="276", after_pt=6)
        closing_run = closing_para.add_run(closing_text)
        set_font(closing_run, size_pt=11)

    # ── Sign-off ──────────────────────────────────────────────────────────────
    sign_para = doc.add_paragraph()
    set_paragraph_spacing(sign_para)
    sign_run = sign_para.add_run(content.get("sign_off", "Sincerely") + ",")
    set_font(sign_run, size_pt=11)

    sig_para = doc.add_paragraph()
    set_paragraph_spacing(sig_para)
    sig_run = sig_para.add_run(profile.get("name", ""))
    set_font(sig_run, size_pt=11, bold=True)

    # ── Save ──────────────────────────────────────────────────────────────────
    doc.save(str(output_path))
    return output_path
